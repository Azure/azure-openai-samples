import docx2txt
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import re
import markdown
import datetime
import chardet

from PySide6.QtCore import Qt
from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QFileDialog, QTextEdit, QToolButton, QSizePolicy, QSplitter
from PySide6.QtGui import QGuiApplication, QTextCursor, QPdfWriter, QTextDocument

import GUI.Common as comm
from GUI.PlainTextPasteEdit import PlainTextPasteEdit
from GUI.ResponseWorker import ResponseWorker

class CompletionsDialog(QDialog):

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Completions")
        self.setGeometry(0, 0, 480, 640)
        self.setWindowFlags(Qt.WindowType.Window)

        self.completions_text = QTextEdit()  
        self.completions_text.setReadOnly(True)
        self.prompt_edit = PlainTextPasteEdit()

        self.commit_button = QPushButton("Commit")
        self.out_button = QPushButton("Save as docs")
        self.in_tool_btn = QToolButton()
        self.in_tool_btn.setText("Load prompt")
        self.in_tool_btn.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Preferred)
        self.in_tool_btn.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.in_text_button = self.in_tool_btn.addAction("Import text")
        self.in_image_button = self.in_tool_btn.addAction("Import image")
        self.out_button.clicked.connect(self.on_export_completions)
        self.commit_button.clicked.connect(self.on_send_prompt)
        self.in_text_button.triggered.connect(self.on_import_prompt)
        self.in_image_button.triggered.connect(self.on_import_image)

        self.in_image_button.setEnabled(comm.vision_flag)

        layout = QVBoxLayout()
        self.splitter = QSplitter(Qt.Orientation.Vertical)
        self.splitter.addWidget(self.completions_text)
        self.splitter.addWidget(self.prompt_edit)
        self.splitter.setSizes([400, 150])
        layout.addWidget(self.splitter)
        
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.in_tool_btn)
        button_layout.addWidget(self.out_button)
        button_layout.addWidget(self.commit_button)
        
        layout.addLayout(button_layout)
        self.setLayout(layout)

        screen = QGuiApplication.primaryScreen()
        screenGeometry = screen.geometry()
        x = (screenGeometry.width() - self.width()) / 2
        y = ((screenGeometry.height() - self.height()) / 2) - 50
        self.move(int(x), int(y))

        self.worker = None
        self.client = parent.client
        self.debug_mode = parent.debug_mode
        self.prompt_template = parent.cfg_prompt_template
        return
      

    def on_send_prompt(self):
        img_paths = self.prompt_edit.getImagePaths()
        user_message = user_message = self.prompt_edit.toPlainText().strip().replace("￼", "")

        if user_message or len(img_paths) > 0:
            if self.debug_mode:
                flag, req = self.client.get_completions_response(user_message, self.prompt_template, img_paths=img_paths)
                self.on_update_completions((flag, req))
            else:
                self.worker = ResponseWorker(self.client, user_message, comps=True, cfg_prompt_template=self.prompt_template, img_paths=img_paths)
                self.worker.signal.connect(self.on_update_completions)
                self.worker.start()

                self.commit_button.setText("Task in progress...")
                self.commit_button.setEnabled(False)
        return
    

    def on_import_prompt(self):
        file_dialog = QFileDialog()
        file_name, _ = file_dialog.getOpenFileName(self, "Import text", "", "Documents (*.txt *.docx *.pdf *.xlsx)")

        if file_name:
            if file_name.endswith('.txt'):
                with open(file_name, "rb") as file:
                    raw_data = file.read()
                    result = chardet.detect(raw_data)
                    encoding = result['encoding']
                    if encoding in ['utf-8', 'UTF-8-SIG']:
                        text = raw_data.decode('utf-8')
                    else:
                        text = raw_data.decode(errors='ignore')
                self.prompt_edit.setPlainText(text)
            elif file_name.endswith('.docx'):
                text = docx2txt.process(file_name)
                self.prompt_edit.setPlainText(text)
            elif file_name.endswith('.pdf'):
                pdf_reader = PdfReader(file_name)
                text = ' '.join([page.extract_text() for page in pdf_reader.pages])
                self.prompt_edit.setPlainText(text)
            elif file_name.endswith('.xlsx'):
                df = pd.read_excel(file_name)
                text = df.to_string(index=False)
                self.prompt_edit.setPlainText(text)
        return
    

    def on_import_image(self):
        file_dialog = QFileDialog()
        file_name, _ = file_dialog.getOpenFileName(self, "Import image", "", "Images (*.jfif *.gif *.jpeg *.jpg *.png *.bmp *.pjp *.apng *.pjpeg *.avif)")

        if file_name:
            with open(file_name, "rb"):
                self.prompt_edit.insertImage(file_name)
        return
    

    def on_update_completions(self, response):
        self.completions_text.clear()
        self.completions_text.moveCursor(QTextCursor.MoveOperation.End)
        flag, req = response
        if flag:
            message = re.sub(r'```(\w*\n)?(.*?)```', comm.replace_code_block, req, flags=re.DOTALL)
            message = markdown.markdown(message)
            self.completions_text.insertHtml(message)
        else:
            self.completions_text.append(req)

        self.commit_button.setText("Commit")
        self.commit_button.setEnabled(True)
        return
    

    def on_export_completions(self):
        document = QTextDocument()
        document.setHtml(self.completions_text.toHtml())
        file_dialog = QFileDialog()
        file_name, _ = file_dialog.getSaveFileName(self, "Save as docs", f"completion_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}", "PDF Files (*.pdf);;Text Files (*.txt);;Word Files (*.docx)")
        if file_name:
            if file_name.endswith('.pdf'):
                pdf_writer = QPdfWriter(file_name)
                document.print_(pdf_writer)
            elif file_name.endswith('.txt'):
                with open(file_name, 'w', encoding='utf-8') as file:
                    file.write(document.toPlainText())
            elif file_name.endswith('.docx'):
                doc = Document()
                doc.add_paragraph(document.toPlainText())
                doc.save(file_name)
        return
    